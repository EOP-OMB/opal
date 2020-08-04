import logging
import django

import docx

from scripts.usefullFunctions import *


def cleanUpRolesAfterImport():
    changeRoll('Information system', 'Information System')
    changeRoll('CIO', 'System Owner')
    changeRoll('Technology Services', 'System Administrators')
    changeRoll('IA', 'Information Assurance Team')
    changeRoll('Technology Services Team Lead', 'Operations Team Lead')
    changeRoll('Administrators', 'System Administrators')
    changeRoll('Information Assurance Team (IAT)', 'Information Assurance Team')
    changeRoll('security engineers', 'Information Assurance Team')
    changeRoll('All', 'organization')
    changeRoll('Organization’s designated personnel or management', 'organization')
    changeRoll('System owners', 'system owner')
    changeRoll('developers', 'Developers')
    changeRoll('MAX.gov systems', 'Information System')


def listDataTables(d):
    for para in d.paragraphs:
        if para.text[0:5] == 'Table':
            print(para.text)


def listTables(document):
    d = docx.Document(document)
    for table in d.tables:
        row1 = []
        for cell in table.row_cells(0):
            row1.append(cell.text)
        print('|' + '|'.join(row1))


def getCheckedOptions(cell):
    results = []
    cell_elm = cell._element
    # checkBoxes = cell_elm.xpath('.//w14:checkBox')
    checkBoxes = cell_elm.xpath(".//*[local-name()='checked']")
    labels = cell_elm.xpath(".//*[local-name()='t']")
    c = 0
    for item in labels:
        if type(item.text) == str and len(item.text) > 1:
            if item.text != "Implementation Status (check all that apply):" \
                    and item.text != "Control Origination (check all that apply):" \
                    and c < len(checkBoxes):
                if checkBoxes[c].values()[0] == '1':
                    results.append(item.text)
                c = c + 1
    return ','.join(results)


def cleanData():
    system_control.objects.all().delete()
    # prop.objects.all().delete()
    # control_parameter.objects.all().delete()
    # control_statement.objects.all().delete()
    # user_role.objects.all().delete()

    # from django.db import connection
    #
    # reset_countersSQL = "delete from sqlite_sequence where name in ('system_control','property','control_parameter','control_statement','user_role');"
    #
    # with connection.cursor() as cursor:
    #     cursor.execute(reset_countersSQL)


def run():
    try:
        cleanData()
        startLogging()
        document = '/Users/dan/PycharmProjects/OSCALweb/Documents/MAX.gov System Security Plan Controls.docx'
        d = docx.Document(document)
        logging.debug('Starting import from ' + document)
        for table in d.tables:
            firstRow = True
            # There are 2 kinds of tables, Control Summary and solution
            # Sometimes the solution table has only 1 column and sometimes it has 2
            columns = len(table.rows[0].cells)
            if table.rows[0].cells[0].text != table.rows[0].cells[columns - 1].text:
                if table.rows[0].cells[1].text == "Control Summary Information" or table.rows[0].cells[
                    1].text == "Control Enhancement Summary Information":
                    logging.debug('Starting import for ' + table.rows[0].cells[0].text + ' control')
                    newControl = system_control.objects.create()
                    newControl.title = table.rows[0].cells[0].text
                    newControl.short_name = str(table.rows[0].cells[0].text).replace('(','.').replace(')','').replace(' ','').lower()
                    newControl.save()
                    logging.debug('Control ' + table.rows[0].cells[0].text + ' created')
                    # Find and Add user_role
                    roleText = table.rows[1].cells[0].text
                    roleList = (roleText[roleText.find(':') + 1:]).split(',')
                    for item in roleList:
                        if item.find('and') > 0:
                            sub_items = item.split('and')
                            for sub_item in sub_items:
                                roleList.append(sub_item)
                            break
                        item = item.strip()
                        logging.debug('Adding role ' + item)
                        newRole, created = user_role.objects.get_or_create(title=item[:100], short_name=item[:25],desc=item[:100])
                        if created:
                            newRole.save()
                        newControl.control_responsible_roles.add(newRole.pk)
                    # Find and Add Parameters, Implementation Status, and Control Origination
                    for row in table.rows:
                        content = row.cells[0].text
                        if content[0:9] == "Parameter":
                            logging.debug('Adding parameter ' + content)
                            newParameter, created = control_parameter.objects.get_or_create(
                                control_parameter_id=content[10:content.find(':')].strip()[:25],
                                value=content[content.find(':') + 1:].strip())
                            if created:
                                newParameter.save()
                            newControl.control_parameters.add(newParameter.pk)
                        elif content[0:14] == "Implementation":
                            logging.debug('*****DEBUG*****' + content)
                            content = getCheckedOptions(row.cells[0])
                            logging.debug('Adding Implementation Status ' + content)
                            newControl.control_status = content
                        elif content[0:7] == "Control":
                            content = getCheckedOptions(row.cells[0])
                            logging.debug('Adding Control Origination ' + content)
                            newControl.control_origination = content
                    newControl.save()
                else:
                    msg = '|'
                    for item in table.rows[0].cells:
                        msg = msg + item.text + '|'
                    logging.debug('Table not imported. First Row text was ' + msg)
            else:
                rowCount = 0
                conID = ''
                for row in table.rows:
                    # skip header row
                    if firstRow:
                        firstRow = False
                        t = row.cells[0].text
                        conID = t[0:t.find(' What')]
                    else:
                        logging.debug('Importing control_statement for ' + conID)
                        control = system_control.objects.get(title=conID)
                        # Handle the case where the table has only 1 column
                        if len(row.cells) == 1:
                            controlName = 'Part a'
                            controlValue = row.cells[0].text
                        else:
                            controlName = row.cells[0].text
                            controlValue = row.cells[1].text
                        logging.debug('Adding statement ' + controlName + ': ' + controlValue)
                        ctrlStatement = control_statement.objects.create(title=controlName, control_statement_text=controlValue)
                        control.control_statements.add(ctrlStatement.pk)
                        rowCount = rowCount + 1
    except django.db.utils.IntegrityError:
        print(uuid)


# import docx
# from ssp import models
# import logging
# from scripts.importSecurityControlsFromWord import getCheckedOptions, main, cleanData, listRolesWithControlCount, changeRoll, delUnusedRoles
# f = '/Users/dan/PycharmProjects/OSCALweb/Documents/MAX.gov System Security Plan Controls.docx'
